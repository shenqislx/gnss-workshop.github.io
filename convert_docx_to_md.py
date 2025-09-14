#!/usr/bin/env python3
"""
DOCX to Markdown转换脚本
用于将Word文档转换为Markdown格式
"""

import sys
import os
from docx import Document
import re

def docx_to_markdown(docx_path, md_path):
    """将DOCX文件转换为Markdown格式"""
    try:
        # 读取DOCX文件
        doc = Document(docx_path)
        
        # 准备Markdown内容
        md_content = []
        
        # 添加元数据
        md_content.append("---")
        md_content.append(f"title: \"粗时导航五状态方程\"")
        md_content.append("date: 2024-01-25")
        md_content.append("author: \"Andy\"")
        md_content.append("category: \"GNSS定位算法\"")
        md_content.append("---")
        md_content.append("")
        
        # 处理文档内容
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # 简单的格式转换
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name[7:]) if len(paragraph.style.name) > 7 else 1
                    md_content.append(f"{'#' * level} {text}")
                else:
                    md_content.append(text)
                md_content.append("")
        
        # 处理表格
        for table in doc.tables:
            md_content.append("")
            for i, row in enumerate(table.rows):
                cells = [cell.text for cell in row.cells]
                if i == 0:  # 表头
                    md_content.append("| " + " | ".join(cells) + " |")
                    md_content.append("|" + " --- |" * len(cells))
                else:
                    md_content.append("| " + " | ".join(cells) + " |")
            md_content.append("")
        
        # 写入Markdown文件
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(md_content))
        
        print(f"成功转换: {docx_path} -> {md_path}")
        return True
        
    except Exception as e:
        print(f"转换失败: {e}")
        return False

def main():
    # 输入和输出文件路径
    docx_file = "articles/粗时导航五状态方程.docx"
    md_file = "articles/coarse-time-navigation-five-state.md"
    
    # 检查文件是否存在
    if not os.path.exists(docx_file):
        print(f"文件不存在: {docx_file}")
        return
    
    # 转换文件
    success = docx_to_markdown(docx_file, md_file)
    
    if success:
        print("转换完成！")
    else:
        print("请确保已安装python-docx库: pip install python-docx")

if __name__ == "__main__":
    main()